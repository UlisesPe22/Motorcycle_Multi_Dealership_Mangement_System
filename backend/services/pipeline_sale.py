import os
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from sqlalchemy.orm import Session
from num2words import num2words

from config import (
    STORAGE_ROOT, CONTRACT_TEMPLATE_PATH,
    SOLICITUD_TEMPLATE_PATH, DISCOUNTS_ACTIVE,
    CONTRACTS_STORAGE_PATH,
)
def format_price_text(price: float) -> str:
    pesos = int(price)
    centavos = round((price - pesos) * 100)
    words = num2words(pesos, lang='es').upper()
    return f"{words} PESOS {centavos:02d}/100 M.N."

def get_moto_price(catalog_entry) -> float:
    if DISCOUNTS_ACTIVE:
        return catalog_entry.discount_price
    return catalog_entry.full_price


def format_price(price: float) -> str:
    return f"${price:,.2f}"


def generate_contract_number(db: Session, dealership_id: int) -> str:
    from models.dealership import Dealership
    from models.contract import Contract
    dealership = db.query(Dealership).filter(
        Dealership.dealership_id == dealership_id
    ).first()
    count = db.query(Contract).filter(
        Contract.dealership_id == dealership_id
    ).count()
    return f"{dealership.contract_prefix}{count + 1:06d}"


def replace_placeholders(doc: Document, replacements: dict) -> None:
    """
    Replaces all {{ KEY }} placeholders in DOCX.
    Handles split runs by rebuilding paragraph text.
    Iterates both paragraphs and all table cells.
    Placeholder format in document: {{ KEY }}
    with spaces inside the double braces.
    """
    def replace_in_paragraph(para):
        for key, value in replacements.items():
            placeholder = f"{{{{ {key} }}}}"
            if placeholder in para.text:
                full_text = para.text.replace(
                    placeholder, str(value or ""))
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = full_text

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)


def build_contract_replacements(contract, db: Session) -> dict:
    from models.motorcycle_model_code import MotorcycleModelCode

    client     = contract.client
    moto       = contract.motorcycle
    catalog    = moto.model
    dealership = contract.dealership
    employee   = contract.employee
    price      = get_moto_price(catalog)

    dt = contract.created_at
    months_es = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre",
        12: "diciembre"
    }

    code_entry = db.query(MotorcycleModelCode).filter(
        MotorcycleModelCode.model_id == catalog.model_id
    ).first()
    moto_code = code_entry.modelo_code if code_entry else ""

    return {
        "CONTRACT_NUMBER":    contract.contract_number,
        "SELLER_OWNER_NAME":  employee.name,
        "DEALERSHIP_NAME":    dealership.name_contract or "",
        "BUYER_NAME":         client.nombre_completo or "",
        "DEALERSHIP_ADDRESS": dealership.address or "",
        "DEALERSHIP_PHONE":   "",
        "DEALERSHIP_EMAIL":   "",
        "DEALERSHIP_RFC":     "",
        "SALE_DATE":          dt.strftime("%Y-%m-%d"),
        "SALE_DATETIME":      dt.strftime("%Y-%m-%d %H:%M:%S"),
        "SIGNATURE_CITY":     dealership.city_contract or "",
        "SIGNATURE_DAY":      str(dt.day),
        "SIGNATURE_MONTH":    months_es[dt.month],
        "SIGNATURE_YEAR":     str(dt.year),
        "BUYER_ADDRESS":      client.domicilio or "",
        "BUYER_RFC":          (client.curp[:13] if client.curp else ""),
        "BUYER_PHONE":        getattr(client, "phone", None) or "",
        "BUYER_EMAIL":        getattr(client, "email", None) or "",
        "MOTO_SERIE":         moto.reference_number or "",
        "MOTO_MOTOR":         moto.motor_number or "",
        "MOTO_MODEL":         catalog.canonical_name or "",
        "MOTO_CODE":          moto_code,
        "MOTO_YEAR":          catalog.year or "",
        "MOTO_COLOR":         moto.color or "",
        "MOTO_BRAND":         "BAJAJ",
        "MOTO_COUNTRY":       "India",
        "MOTO_CAPACITY":      "2 PASAJEROS",
        "MOTO_PEDIMENTO":     "25 51 1626 5004107",
        "MOTO_ADUANA":        "510-LAZARO CARDENAS, MICHOACAN",
        "MOTO_PRICE_NUMBER":  format_price(price),
        "MOTO_PRICE_TEXT": format_price_text(price),
        "PAYMENT_AMOUNT":     format_price(
                                  contract.payment_downpayment or price),
        "PAYMENT_CONCEPT":    contract.sale_type.value.upper(),
        "PAYMENT_METHOD":     (contract.payment_method.value.upper()
                               if contract.payment_method else ""),
    }


def build_solicitud_replacements(contract, db: Session) -> dict:
    base  = build_contract_replacements(contract, db)
    price = get_moto_price(contract.motorcycle.model)
    base.update({
        "PAYMENT_DOWNPAYMENT":     format_price(
            contract.payment_downpayment or 0),
        "PAYMENT_PENDING":         format_price(
            price - (contract.payment_downpayment or 0)),
        "PAYMENT_BANK":            contract.payment_bank or "",
        "PAYMENT_FINANCE_COMPANY": (contract.institution.name
                                    if contract.institution else ""),
        "PAYMENT_CREDIT_TYPE":     "CREDITO",
        "REFERENCE_NAME":          contract.reference_name or "",
        "REFERENCE_PHONE":         contract.reference_phone or "",
        "REFERENCE_RELATION":      contract.reference_relation or "",
        "BUYER_COLONIA":           contract.buyer_colonia or "",
        "BUYER_CP":                contract.buyer_cp or "",
        "BUYER_MUNICIPIO":         contract.buyer_municipio or "",
        "BUYER_ESTADO":            contract.buyer_estado or "",
    })
    return base


def generate_documents(contract, db: Session) -> tuple[str, Optional[str]]:
    os.makedirs(CONTRACTS_STORAGE_PATH, exist_ok=True)

    contract_docx_path = os.path.join(
        CONTRACTS_STORAGE_PATH,
        f"{contract.contract_number}_contrato.docx"
    )
    doc = Document(CONTRACT_TEMPLATE_PATH)
    replace_placeholders(doc, build_contract_replacements(contract, db))
    doc.save(contract_docx_path)

    solicitud_docx_path = None
    if contract.sale_type.value == "credito":
        solicitud_docx_path = os.path.join(
            CONTRACTS_STORAGE_PATH,
            f"{contract.contract_number}_solicitud.docx"
        )
        sol_doc = Document(SOLICITUD_TEMPLATE_PATH)
        replace_placeholders(
            sol_doc, build_solicitud_replacements(contract, db))
        sol_doc.save(solicitud_docx_path)

    return contract_docx_path, solicitud_docx_path
